import pandas as pd 
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
worddic={'身分證號碼':'身分證號',
         '中文姓名':'姓名',
         '出生日期':'出生日期',
         '英文姓名':'英文姓名',
         '原住民傳統姓名並列之羅馬拼音':'身分證上原住民姓名之羅馬拼音',
         '通訊地址':'通訊地址',
         '戶籍地址':'戶籍地址',
         '就讀學校':'報檢人參檢學校',
         '就讀科系':'科系',
         '年級':'年級',
         '班別':'班別',
         '上課別':'部別'
         }
#讀取主要資料
df = pd.read_excel('1.中壢高商(14901).xlsx', sheet_name='Data-全測',skiprows=2)
rows = df.shape[0]
#--------------------------------------------------------------------------------
#讀取&儲存科系資料
dfs = pd.read_excel('1.中壢高商(14901).xlsx',sheet_name='代號',usecols=[0])
text_values = dfs.values
is_string =np.vectorize(lambda x: isinstance(x, str))(text_values).astype(bool)
only_string = text_values[is_string]
subject = np.insert(only_string,0,0)
#---------------------------------------------------------------------------------
#讀取&儲存班級資料
dfc  = pd.read_excel('1.中壢高商(14901).xlsx',sheet_name='代號',usecols=[3])
text_values = dfc.values
is_string =np.vectorize(lambda x: isinstance(x, str))(text_values).astype(bool)
only_string = text_values[is_string]
Class = np.insert(only_string,0,0)
#---------------------------------------------------------------------------------
#報檢學校編號與代碼
dfsc = pd.read_excel('1.中壢高商(14901).xlsx',sheet_name='代號',usecols=[7,8,9], index_col=0,nrows = 17)
#---------------------------------------------------------------------------------
#學制
dfstu  = pd.read_excel('1.中壢高商(14901).xlsx',sheet_name='代號',usecols=[18],nrows = 4)
text_values = dfstu.values
is_string =np.vectorize(lambda x: isinstance(x, str))(text_values).astype(bool)
only_string = text_values[is_string]
stu = np.insert(only_string,0,0)
#---------------------------------------------------------------------------------
#特定對象
dfsg = pd.read_excel('1.中壢高商(14901).xlsx',sheet_name='代號',usecols=[20,21],index_col=0, nrows = 8)
#---------------------------------------------------------------------------------
#測驗類別
dftp = pd.read_excel('1.中壢高商(14901).xlsx',sheet_name='代號',usecols=[14,15],index_col=0, nrows = 13 )
text_values = dftp.values
is_string =np.vectorize(lambda x: isinstance(x, str))(text_values).astype(bool)
only_string = text_values[is_string]
test_type_lst = np.insert(only_string,0,0)
#---------------------------------------------------------------------------------
#套印用資料-全測
df_print = pd.read_excel('1.中壢高商(14901).xlsx', sheet_name='套印用資料-全測')
#---------------------------------------------------------------------------------
#word讀取 & 填寫
for i in range(0,rows):
    doc = Document('5.報名表正面.docx')
    table =doc.tables[0]
    nowcommend = ''
    testset = set()
    for idxr,row in enumerate(table.rows) :
        for idxc, cell in enumerate(row.cells) :
            if cell.text != '':
                if cell.text not in testset :
                    if nowcommend=='英文姓名':
                        if cell.text not in testset:
                            care = cell.text
                            paragraph = cell.paragraphs[0]
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.line_spacing = Pt(12)
                            for run in paragraph.runs:
                                run.clear()
                            #---------------------------------------------------------------------------------
                            run2 = paragraph.add_run(df.loc[i, worddic[nowcommend]])  
                            run2.font.size = Pt(12)  # 設置字體大小 
                            run2.font.name = 'Times New Roman'
                            run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                            #---------------------------------------------------------------------------------
                            run1 = paragraph.add_run(care)  
                            run1.font.size = Pt(6)
                            run1.font.name = '標楷體'
                            run1.font.bold = True
                            run1._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                            testset.add(df.loc[i,worddic[nowcommend]] +care)
                    elif nowcommend == '聯絡電話':
                         if cell.text not in testset:
                            homnnum,clphone = cell.text.split('\n')
                            homnnum = homnnum.replace(' ','')
                            clphone = clphone.replace(' ','')
                            homnnum = homnnum + str(df.loc[i,'電話(公)'])
                            clphone = clphone + '0' +str(df.loc[i,'電話(行動)'])
                            cell.text = homnnum + '\n' + clphone
                            testset.add(cell.text)
                    elif nowcommend =='上課別':
                        index = df.loc[i,str(worddic[nowcommend])]
                        reCell  = table.cell(idxr+1 , idxc-2) 
                        reCell.text  = stu[int(index)]
                    elif nowcommend =='年級':
                        reCell  = table.cell(idxr+1 , idxc-1)
                        reCell.text = str(df.loc[i,'年級'])
                    elif nowcommend =='班別':
                        reCell  = table.cell(idxr+1 , idxc-1)
                        reCell.text = str(df.loc[i, '班別'])
                    elif nowcommend == '身分別(一般報檢人免填)':
                        idCode =str(df.loc[i,'特定對象身份別'])
                        if idCode != 'nan':
                            id = dfsg.loc[float(idCode), 'Unnamed: 21']
                            checkboxlst = list(map(str,cell.text.split('\n')))
                            cell.text = ''
                            option_lst = [] 
                            paragraph = cell.paragraphs[0]
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.line_spacing = Pt(9)
                            underline_option = ''
                            for j in checkboxlst:
                                times = j.count('□')
                                temlst = []
                                if times == 0 :
                                    underline_option +=j +'\n'
                                    option_lst.append([j,'\n'])
                                else:
                                    for k in range(times):
                                        option_not_edit = j[len(j)-j[::-1].index('□')-1:]
                                        squard , option = option_not_edit[:1],option_not_edit[1:]
                                        if k == 0 :
                                            option +='\n'
                                        temlst.append([squard,option])
                                        j = j.replace(option_not_edit,'')
                                temlst = temlst[::-1]
                                for x in temlst :
                                    option_lst.append(x)
                            for option in option_lst:
                                if id in option[1]:
                                    option[0] = '■'
                                    break
                            response = ''
                            for res in option_lst:
                                run = paragraph.add_run(res[0])
                                run.font.size = Pt(7.8)
                                run.font.name = '標楷體'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                                run2 = paragraph.add_run(res[1])
                                run2.font.size = Pt(7.8)
                                run2.font.name ='Times New Roman'  
                                run2._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')     
                                if '限以上身分，需另填寫補助申請書，不申請補助者免附' in res[0]:
                                    run3 = paragraph.add_run('————————————————————\n') 
                                    run3.font.name ='Times New Roman'  
                                    run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')             
                                    rPr2 = run3._element.get_or_add_rPr()
                                    spacing = OxmlElement('w:spacing')
                                    spacing.set(qn('w:val'), '0.5')  
                                    rPr2.append(spacing)
                                rPr = run2._element.get_or_add_rPr()
                                spacing = OxmlElement('w:spacing')
                                spacing.set(qn('w:val'), '0.5')  
                                rPr.append(spacing)
                    elif nowcommend =='報檢職類':
                        id_type = int(df_print.loc[i,'測驗類別'])
                        test_type = test_type_lst[id_type]
                        checkboxlst = list(map(str,cell.text.split('\n')))
                        option_lst = []
                        cell.text = ''
                        paragraph = cell.paragraphs[0]
                       # ---------------------------------------------------------------------------------
                       # 分割checkbox 與選項內容
                        for j in checkboxlst:
                                times = j.count('□')
                                temlst = []
                                if times == 0 :
                                    underline_option +=j +'\n' 
                                    option_lst.append([j,'\n'])
                                else:
                                    for k in range(times):
                                        option_not_edit = j[len(j)-j[::-1].index('□')-1:]
                                        squard , option = option_not_edit[:1],option_not_edit[1:]
                                        if k == 0 :
                                            option +='\n'
                                        temlst.append([squard,option])
                                        j = j.replace(option_not_edit,'')
                                temlst = temlst[::-1]
                                for x in temlst :
                                    option_lst.append(x)
                        for res in option_lst:
                            if test_type in res[1]:
                                res[0] = '■'
                                break
                        for res in option_lst:
                                run = paragraph.add_run(res[0])
                                run.font.size = Pt(12)
                                run.font.name = '標楷體'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                                run2 = paragraph.add_run(res[1])
                                run2.font.size = Pt(12)
                                run2.font.name ='Arial'  
                                run2._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')  
                                rPr = run._element.get_or_add_rPr()
                                spacing = OxmlElement('w:spacing')
                                spacing.set(qn('w:val'), '0')  
                                rPr.append(spacing)
                                rPr = run2._element.get_or_add_rPr()
                                spacing = OxmlElement('w:spacing')
                                spacing.set(qn('w:val'), '0')  
                                rPr.append(spacing) 
                                tr = row._element  # 获取表格行的 XML 元素
                                # trPr = tr.get_or_add_trPr()  # 获取或创建 trPr 元素
                                # trHeight = OxmlElement('w:trHeight')  # 创建行高元素
                                # trHeight.set('w:val', '400')  # 设置行高，单位是1/20磅，400表示20磅的行高
                                # trHeight.set('w:hRule', 'exact')  # 设置为 exact 表示固定行高
                                # trPr.append(trHeight)  # 添加到行属性中  
                       # ---------------------------------------------------------------------------------
                    # elif '實貼身分證【正面】' in nowcommend:
                    testset.add(cell.text)
                    nowcommend = cell.text
                    nowcommend = nowcommend.replace('\n','')
            else:
                    try:
                        if nowcommend =='就讀學校':
                            school  = df.loc[i,worddic[nowcommend]]
                            cell.text = dfsc.loc[str(school),'Unnamed: 9']
                        elif nowcommend =='就讀科系':
                            sub = df.loc[i,worddic[nowcommend]]
                            cell.text=subject[int(sub)]
                        else:
                            report = str(df.loc[i,worddic[nowcommend]])
                            if report =='nan':
                                cell.text = ''
                            else:
                                cell.text = report 
                    except:
                         break
    new_file_path = str(i+1)+'.docx'
    doc.save(new_file_path) 